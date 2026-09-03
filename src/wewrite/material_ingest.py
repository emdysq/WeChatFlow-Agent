"""Deterministic local material ingestion for one-command composition."""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from pathlib import Path

import yaml
from bs4 import BeautifulSoup
from PIL import Image


TEXT_SUFFIXES = {".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".html", ".htm"}
DOCUMENT_SUFFIXES = {".pdf", ".docx"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | DOCUMENT_SUFFIXES | IMAGE_SUFFIXES


@dataclass
class TextMaterial:
    material_id: str
    name: str
    source_path: str
    content: str
    truncated: bool = False


@dataclass
class ImageMaterial:
    asset_id: str
    name: str
    source_path: str
    relative_path: str
    width: int | None = None
    height: int | None = None
    analysis: dict = field(default_factory=dict)


@dataclass
class MaterialPack:
    texts: list[TextMaterial]
    images: list[ImageMaterial]
    skipped: list[str] = field(default_factory=list)

    def as_prompt(self) -> str:
        parts = ["# 用户素材包", "", "以下内容全部来自用户本次提供的材料。"]
        for item in self.texts:
            suffix = "（内容过长，已截取）" if item.truncated else ""
            parts += ["", f"## [{item.material_id}] {item.name}{suffix}", "", item.content]
        for image in self.images:
            analysis = image.analysis or {"status": "unparsed"}
            parts += [
                "",
                f"## [{image.asset_id}] 图片：{image.name}",
                "",
                yaml.safe_dump(analysis, allow_unicode=True, sort_keys=False).strip(),
            ]
        return "\n".join(parts).strip() + "\n"


def _read_text_file(path: Path) -> str:
    raw = path.read_text(encoding="utf-8-sig", errors="replace")
    if path.suffix.lower() in {".html", ".htm"}:
        return BeautifulSoup(raw, "html.parser").get_text("\n", strip=True)
    return raw


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("读取 PDF 需要安装 pypdf") from exc
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("读取 DOCX 需要安装 python-docx") from exc
    document = Document(str(path))
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                blocks.append(" | ".join(values))
    return "\n".join(blocks)


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _read_text_file(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"不支持的素材格式: {path}")


def _expand_paths(values: list[str | Path]) -> tuple[list[Path], list[str]]:
    files: list[Path] = []
    skipped: list[str] = []
    for value in values:
        path = Path(value).expanduser()
        if not path.exists():
            raise FileNotFoundError(f"素材不存在: {path}")
        candidates = [path] if path.is_file() else sorted(
            item for item in path.rglob("*")
            if item.is_file() and not any(part.startswith(".") for part in item.relative_to(path).parts)
        )
        for item in candidates:
            if item.suffix.lower() in SUPPORTED_SUFFIXES:
                files.append(item.resolve())
            else:
                skipped.append(str(item))
    unique: list[Path] = []
    seen: set[str] = set()
    for path in files:
        key = str(path).casefold()
        if key not in seen:
            seen.add(key)
            unique.append(path)
    return unique, skipped


def _copy_image(path: Path, assets_dir: Path, index: int) -> ImageMaterial:
    assets_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in path.stem if c.isalnum() or c in {"-", "_"}).strip("._")
    safe_name = safe_name or f"image-{index}"
    target = assets_dir / f"{index:02d}-{safe_name}{path.suffix.lower()}"
    shutil.copy2(path, target)
    width = height = None
    try:
        with Image.open(target) as image:
            width, height = image.size
    except OSError:
        pass
    return ImageMaterial(
        asset_id=f"IMG{index}",
        name=path.name,
        source_path=str(path),
        relative_path=str(Path("assets") / target.name).replace("\\", "/"),
        width=width,
        height=height,
    )


def ingest_materials(
    *,
    material_paths: list[str | Path],
    image_paths: list[str | Path],
    run_directory: str | Path,
    notes: str = "",
    max_chars_per_file: int = 50_000,
    max_total_chars: int = 180_000,
) -> MaterialPack:
    """Read supported files and copy images into the run without altering sources."""
    material_files, skipped = _expand_paths(material_paths)
    explicit_images, image_skipped = _expand_paths(image_paths)
    skipped.extend(image_skipped)
    all_files = material_files + [path for path in explicit_images if path not in material_files]

    texts: list[TextMaterial] = []
    image_sources: list[Path] = []
    if notes.strip():
        texts.append(TextMaterial("M1", "命令行补充说明", "user-input://notes", notes.strip()))

    remaining = max_total_chars - sum(len(item.content) for item in texts)
    for path in all_files:
        suffix = path.suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            image_sources.append(path)
            continue
        if remaining <= 0:
            skipped.append(f"{path}（超过素材总长度上限）")
            continue
        content = _extract_text(path).strip()
        if not content:
            skipped.append(f"{path}（未提取到文本）")
            continue
        limit = min(max_chars_per_file, remaining)
        truncated = len(content) > limit
        content = content[:limit]
        material_id = f"M{len(texts) + 1}"
        texts.append(TextMaterial(material_id, path.name, str(path), content, truncated))
        remaining -= len(content)

    images: list[ImageMaterial] = []
    assets_dir = Path(run_directory) / "assets"
    seen_images: set[str] = set()
    for path in image_sources:
        key = str(path).casefold()
        if key in seen_images:
            continue
        seen_images.add(key)
        images.append(_copy_image(path, assets_dir, len(images) + 1))

    return MaterialPack(texts=texts, images=images, skipped=skipped)
